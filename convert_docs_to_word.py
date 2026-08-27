"""
Comprehensive Word Document (.docx) Generator for Week 4 Project Report
Author: Ajay M - Senior Data Scientist, ML Engineer & Logistics Analytics Consultant
Target: docs/Week4_Predictive_Modeling_and_Optimization.docx
"""

import os
import re
import docx
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import parse_xml, OxmlElement
from docx.oxml.ns import nsdecls, qn

project_root = os.path.dirname(os.path.abspath(__file__))
doc_path = os.path.join(project_root, "docs", "Week4_Predictive_Modeling_and_Optimization.docx")
md_path = os.path.join(project_root, "docs", "Week4_Predictive_Modeling_and_Optimization.md")
figures_dir = os.path.join(project_root, "outputs", "figures")

# Brand Colors
COLOR_NAVY_HEX = "1E3A8A"
COLOR_TEAL_HEX = "0D9488"
COLOR_SLATE_HEX = "0F172A"
COLOR_LIGHT_BG_HEX = "F8FAFC"
COLOR_ROW_ALT_HEX = "F1F5F9"
COLOR_BORDER_HEX = "CBD5E1"

RGB_NAVY = RGBColor(30, 58, 138)
RGB_TEAL = RGBColor(13, 148, 136)
RGB_SLATE = RGBColor(15, 23, 42)
RGB_GRAY = RGBColor(100, 116, 139)
RGB_WHITE = RGBColor(255, 255, 255)


def set_cell_background(cell, fill_hex):
    """Sets background color of a table cell."""
    tcPr = cell._element.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)


def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    """Sets internal padding for a table cell."""
    tcPr = cell._element.get_or_add_tcPr()
    tcMar = parse_xml(
        f'<w:tcMar {nsdecls("w")}>'
        f'<w:top w:w="{top}" w:type="dxa"/>'
        f'<w:bottom w:w="{bottom}" w:type="dxa"/>'
        f'<w:left w:w="{left}" w:type="dxa"/>'
        f'<w:right w:w="{right}" w:type="dxa"/>'
        f'</w:tcMar>'
    )
    tcPr.append(tcMar)


def create_callout_box(doc, title, text, border_color_hex=COLOR_TEAL_HEX, bg_hex=COLOR_LIGHT_BG_HEX):
    """Creates a styled callout box in Word."""
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl.autofit = False
    
    cell = tbl.cell(0, 0)
    cell.width = Inches(6.5)
    set_cell_background(cell, bg_hex)
    set_cell_margins(cell, top=120, bottom=120, left=200, right=150)
    
    tcPr = cell._element.get_or_add_tcPr()
    borders = parse_xml(
        f'<w:tcBorders {nsdecls("w")}>'
        f'<w:top w:val="none"/>'
        f'<w:left w:val="single" w:sz="36" w:space="0" w:color="{border_color_hex}"/>'
        f'<w:bottom w:val="none"/>'
        f'<w:right w:val="none"/>'
        f'</w:tcBorders>'
    )
    tcPr.append(borders)
    
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(3)
    run_title = p.add_run(f"📌 {title}\n")
    run_title.font.name = "Arial"
    run_title.font.size = Pt(10.5)
    run_title.font.bold = True
    run_title.font.color.rgb = RGB_NAVY
    
    run_text = p.add_run(text)
    run_text.font.name = "Arial"
    run_text.font.size = Pt(9.5)
    run_text.font.color.rgb = RGB_SLATE
    
    doc.add_paragraph().paragraph_format.space_after = Pt(4)


def build_word_document():
    doc = Document()
    
    # Page setup - Margins
    for section in doc.sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)
        
        # Header / Footer
        header = section.header
        hp = header.paragraphs[0]
        hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        h_run = hp.add_run("Week 4: Predictive Modeling & Optimization | Ajay M")
        h_run.font.name = "Arial"
        h_run.font.size = Pt(8.5)
        h_run.font.color.rgb = RGB_GRAY
        
        footer = section.footer
        fp = footer.paragraphs[0]
        fp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        f_run = fp.add_run("Candidate: Ajay M  •  Track: Advanced Logistics Analytics & ML Internship")
        f_run.font.name = "Arial"
        f_run.font.size = Pt(8.5)
        f_run.font.color.rgb = RGB_GRAY
        
    # --- 1. COVER PAGE ---
    cover_table = doc.add_table(rows=1, cols=1)
    cover_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cover_table.autofit = False
    c_cell = cover_table.cell(0, 0)
    c_cell.width = Inches(6.5)
    set_cell_background(c_cell, COLOR_NAVY_HEX)
    set_cell_margins(c_cell, top=450, bottom=450, left=320, right=320)
    
    cp = c_cell.paragraphs[0]
    cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    tag_run = cp.add_run("ACADEMIC INTERNSHIP & TECHNICAL PROJECT REPORT\n\n")
    tag_run.font.name = "Arial"
    tag_run.font.size = Pt(11)
    tag_run.font.bold = True
    tag_run.font.color.rgb = RGBColor(203, 213, 225)
    
    title_run = cp.add_run("WEEK 4: PREDICTIVE MODELING AND OPTIMIZATION IN LOGISTICS SYSTEMS\n\n")
    title_run.font.name = "Arial"
    title_run.font.size = Pt(20)
    title_run.font.bold = True
    title_run.font.color.rgb = RGB_WHITE
    
    sub_run = cp.add_run("Logistics Delivery Time Prediction, Anti-Leakage Pipeline Benchmarking, and Multi-Region Operational Resource Optimization\n\n\n")
    sub_run.font.name = "Arial"
    sub_run.font.size = Pt(12)
    sub_run.font.color.rgb = RGBColor(226, 232, 240)
    
    meta_box = cp.add_run(
        "━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━\n"
        "CANDIDATE INFORMATION & METADATA\n"
        "━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━\n"
        "Author / Candidate Name: Ajay M\n"
        "Role: Senior Data Scientist, ML Engineer & Optimization Consultant\n"
        "Academic Track: Logistics Analytics & Machine Learning Internship\n"
        "Technical Stack: Python 3.14 / Scikit-Learn 1.9.0 / SciPy 1.18.1\n"
        "GitHub Repository: https://github.com/ajaymari74-cyber/logisticweek4.git\n"
        "Submission Date: August 2026\n"
    )
    meta_box.font.name = "Arial"
    meta_box.font.size = Pt(10)
    meta_box.font.color.rgb = RGB_WHITE
    
    doc.add_page_break()
    
    # --- 2. CANDIDATE DECLARATION & APPROVAL PAGE ---
    h_dec = doc.add_heading("Certificate of Originality & Candidate Declaration", level=1)
    h_dec.paragraph_format.space_before = Pt(10)
    h_dec.paragraph_format.space_after = Pt(12)
    h_dec.runs[0].font.name = "Arial"
    h_dec.runs[0].font.size = Pt(15)
    h_dec.runs[0].font.color.rgb = RGB_NAVY
    
    p_dec1 = doc.add_paragraph()
    p_dec1.paragraph_format.line_spacing = 1.2
    r_dec1 = p_dec1.add_run(
        "I, Ajay M, hereby declare that this project report entitled \"Week 4: Predictive Modeling and Optimization in Logistics Systems\" "
        "is a bonafide record of independent analytical and technical work executed by me under the Advanced Logistics Analytics and Machine Learning Internship Program. "
        "All data preprocessing, feature engineering, machine learning modeling, cross-validation, hyperparameter tuning, operations research optimization, "
        "and empirical result documentation were developed and validated by me using reproducible Python workflows."
    )
    r_dec1.font.name = "Arial"
    r_dec1.font.size = Pt(10)
    r_dec1.font.color.rgb = RGB_SLATE
    
    doc.add_paragraph()
    
    sig_table = doc.add_table(rows=2, cols=2)
    sig_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    sig_table.autofit = False
    sig_table.columns[0].width = Inches(3.2)
    sig_table.columns[1].width = Inches(3.2)
    
    cell_s1 = sig_table.cell(0, 0)
    p_s1 = cell_s1.paragraphs[0]
    p_s1.add_run("Candidate Signature:\n\n_______________________\n").font.bold = True
    r_name = p_s1.add_run("Ajay M\nLead Data Scientist & Candidate")
    r_name.font.color.rgb = RGB_NAVY
    
    cell_s2 = sig_table.cell(0, 1)
    p_s2 = cell_s2.paragraphs[0]
    p_s2.add_run("Internship Evaluator / Mentor:\n\n_______________________\n").font.bold = True
    r_eval = p_s2.add_run("Academic Project Reviewer\nFaculty of Data Science & Logistics")
    r_eval.font.color.rgb = RGB_NAVY
    
    doc.add_page_break()
    
    # Read Markdown Document
    with open(md_path, "r", encoding="utf-8") as f:
        md_text = f.read()
        
    lines = md_text.split("\n")
    in_table = False
    table_lines = []
    
    # Track inserted figures
    inserted_figures = set()
    
    for line in lines:
        stripped = line.strip()
        
        # Check for Table
        if stripped.startswith("|") and stripped.endswith("|"):
            in_table = True
            table_lines.append(stripped)
            continue
        elif in_table:
            # Process table
            if table_lines:
                parsed_rows = []
                for t_row in table_lines:
                    if re.match(r"^\|[\s\-:|]+\|$", t_row):
                        continue
                    cells_data = [c.strip() for c in t_row.split("|")[1:-1]]
                    parsed_rows.append(cells_data)
                    
                if parsed_rows:
                    num_cols = max(len(r) for r in parsed_rows)
                    tbl = doc.add_table(rows=len(parsed_rows), cols=num_cols)
                    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
                    tbl.autofit = False
                    
                    for r_idx, r_data in enumerate(parsed_rows):
                        for c_idx in range(num_cols):
                            cell_text = r_data[c_idx] if c_idx < len(r_data) else ""
                            cell = tbl.cell(r_idx, c_idx)
                            set_cell_margins(cell, top=80, bottom=80, left=100, right=100)
                            
                            p = cell.paragraphs[0]
                            p.paragraph_format.space_before = Pt(1)
                            p.paragraph_format.space_after = Pt(1)
                            run = p.add_run(cell_text)
                            run.font.name = "Arial"
                            
                            if r_idx == 0:
                                set_cell_background(cell, COLOR_NAVY_HEX)
                                run.font.bold = True
                                run.font.size = Pt(9)
                                run.font.color.rgb = RGB_WHITE
                            else:
                                bg_color = COLOR_ROW_ALT_HEX if r_idx % 2 == 1 else "FFFFFF"
                                set_cell_background(cell, bg_color)
                                run.font.size = Pt(8.5)
                                run.font.color.rgb = RGB_SLATE
                                
                    doc.add_paragraph().paragraph_format.space_after = Pt(4)
            in_table = False
            table_lines = []
            
        # Headings
        if stripped.startswith("# "):
            continue
        elif stripped.startswith("## "):
            h_text = stripped[3:].strip()
            h = doc.add_heading(h_text, level=1)
            h.paragraph_format.space_before = Pt(14)
            h.paragraph_format.space_after = Pt(4)
            for r in h.runs:
                r.font.name = "Arial"
                r.font.size = Pt(13)
                r.font.bold = True
                r.font.color.rgb = RGB_NAVY
                
            # Figure insertions based on sections
            if "16. Model Comparison" in h_text or "17. Best Model" in h_text:
                f_path = os.path.join(figures_dir, "04_model_comparison.png")
                if os.path.exists(f_path) and "04" not in inserted_figures:
                    doc.add_paragraph()
                    doc.add_picture(f_path, width=Inches(6.0))
                    p_cap = doc.add_paragraph("Figure 4: Predictive Model Performance Benchmarks (MAE, RMSE, R²)")
                    p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    p_cap.runs[0].font.size = Pt(8.5)
                    p_cap.runs[0].font.italic = True
                    inserted_figures.add("04")
                    
            elif "14. Cross-Validation" in h_text:
                f_path = os.path.join(figures_dir, "05_cross_validation_stability.png")
                if os.path.exists(f_path) and "05" not in inserted_figures:
                    doc.add_paragraph()
                    doc.add_picture(f_path, width=Inches(5.8))
                    p_cap = doc.add_paragraph("Figure 5: 5-Fold Cross-Validation Stability and Generalization Error")
                    p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    p_cap.runs[0].font.size = Pt(8.5)
                    p_cap.runs[0].font.italic = True
                    inserted_figures.add("05")
                    
            elif "15. Hyperparameter Tuning" in h_text:
                f_path = os.path.join(figures_dir, "07_hyperparameter_tuning_impact.png")
                if os.path.exists(f_path) and "07" not in inserted_figures:
                    doc.add_paragraph()
                    doc.add_picture(f_path, width=Inches(5.8))
                    p_cap = doc.add_paragraph("Figure 7: GridSearchCV Hyperparameter Tuning Impact Comparison")
                    p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    p_cap.runs[0].font.size = Pt(8.5)
                    p_cap.runs[0].font.italic = True
                    inserted_figures.add("07")
                    
            elif "18. Prediction Analysis" in h_text:
                f_path1 = os.path.join(figures_dir, "01_actual_vs_predicted.png")
                f_path2 = os.path.join(figures_dir, "02_residual_analysis.png")
                f_path3 = os.path.join(figures_dir, "03_residual_distribution.png")
                if os.path.exists(f_path1) and "01" not in inserted_figures:
                    doc.add_paragraph()
                    doc.add_picture(f_path1, width=Inches(5.5))
                    p_cap = doc.add_paragraph("Figure 1: Actual vs. Predicted Delivery Time Scatter Plot (y = x Line)")
                    p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    p_cap.runs[0].font.size = Pt(8.5)
                    p_cap.runs[0].font.italic = True
                    inserted_figures.add("01")
                if os.path.exists(f_path2) and "02" not in inserted_figures:
                    doc.add_paragraph()
                    doc.add_picture(f_path2, width=Inches(5.5))
                    p_cap = doc.add_paragraph("Figure 2: Residual Error Diagnostic Plot vs. Fitted Values")
                    p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    p_cap.runs[0].font.size = Pt(8.5)
                    p_cap.runs[0].font.italic = True
                    inserted_figures.add("02")
                if os.path.exists(f_path3) and "03" not in inserted_figures:
                    doc.add_paragraph()
                    doc.add_picture(f_path3, width=Inches(5.8))
                    p_cap = doc.add_paragraph("Figure 3: Residual Density and Normal Q-Q Probability Plot")
                    p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    p_cap.runs[0].font.size = Pt(8.5)
                    p_cap.runs[0].font.italic = True
                    inserted_figures.add("03")
                    
            elif "19. Feature Importance" in h_text:
                f_path = os.path.join(figures_dir, "06_feature_importance.png")
                if os.path.exists(f_path) and "06" not in inserted_figures:
                    doc.add_paragraph()
                    doc.add_picture(f_path, width=Inches(5.8))
                    p_cap = doc.add_paragraph("Figure 6: Top Predictive Feature Importances (Gradient Boosting)")
                    p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    p_cap.runs[0].font.size = Pt(8.5)
                    p_cap.runs[0].font.italic = True
                    inserted_figures.add("06")
                    
            elif "23. Shipping Mode Optimization" in h_text or "24. Business Insights" in h_text:
                f_path1 = os.path.join(figures_dir, "08_optimization_cost_time_tradeoff.png")
                f_path2 = os.path.join(figures_dir, "09_regional_optimization_summary.png")
                if os.path.exists(f_path1) and "08" not in inserted_figures:
                    doc.add_paragraph()
                    doc.add_picture(f_path1, width=Inches(5.8))
                    p_cap = doc.add_paragraph("Figure 8: Cost vs. Delivery Time Pareto Trade-Off Curve")
                    p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    p_cap.runs[0].font.size = Pt(8.5)
                    p_cap.runs[0].font.italic = True
                    inserted_figures.add("08")
                if os.path.exists(f_path2) and "09" not in inserted_figures:
                    doc.add_paragraph()
                    doc.add_picture(f_path2, width=Inches(5.8))
                    p_cap = doc.add_paragraph("Figure 9: Regional Logistics Cost: Baseline vs. Optimized Dispatch")
                    p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    p_cap.runs[0].font.size = Pt(8.5)
                    p_cap.runs[0].font.italic = True
                    inserted_figures.add("09")
            continue
            
        elif stripped.startswith("### "):
            h_text = stripped[4:].strip()
            h = doc.add_heading(h_text, level=2)
            h.paragraph_format.space_before = Pt(10)
            h.paragraph_format.space_after = Pt(2)
            for r in h.runs:
                r.font.name = "Arial"
                r.font.size = Pt(11)
                r.font.bold = True
                r.font.color.rgb = RGB_TEAL
            continue
            
        # Callout alerts
        if stripped.startswith("> [!IMPORTANT]"):
            create_callout_box(doc, "CRITICAL DATA LEAKAGE PROTOCOL", "All post-event features (Shipping_Delay_Days, Is_Delayed, Delivery_Status, Customer_Rating, Speed_Index_KMPD) are strictly excluded from the training matrix.", border_color_hex="E11D48")
            continue
        elif stripped.startswith("> [!NOTE]"):
            create_callout_box(doc, "STATISTICAL INTERPRETATION NOTE", "Feature importance metrics confirm empirical predictive association within the dataset; they do not independently establish physical causation.", border_color_hex=COLOR_TEAL_HEX)
            continue
        elif stripped.startswith(">"):
            continue
            
        # Regular paragraph
        if stripped:
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(3)
            p.paragraph_format.line_spacing = 1.15
            
            parts = re.split(r"(\*\*.*?\*\*)", stripped)
            for pt in parts:
                if pt.startswith("**") and pt.endswith("**"):
                    r = p.add_run(pt[2:-2])
                    r.font.bold = True
                else:
                    r = p.add_run(pt)
                r.font.name = "Arial"
                r.font.size = Pt(9.5)
                r.font.color.rgb = RGB_SLATE
                
    doc.save(doc_path)
    file_size = os.path.getsize(doc_path)
    print(f"Professional Word document compiled successfully: {doc_path} ({file_size / (1024*1024):.2f} MB)")


if __name__ == "__main__":
    build_word_document()
